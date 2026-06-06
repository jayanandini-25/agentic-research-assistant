"""
DOCX Exporter — Phase 11
=====================================================================
File location:  app/export/docx_exporter.py

Converts the markdown research report into a styled Microsoft Word document.

Uses python-docx to produce a properly formatted .docx with:
  - Title style for the report heading (# h1)
  - Heading 1 / Heading 2 styles for sections (## / ###)
  - Normal paragraph style for body text with justified alignment
  - Inline citation numbers [N] formatted as superscript in blue
  - Images embedded inline at correct section positions (downloaded from URLs)
  - Bullet and numbered lists
  - Code blocks with monospace font and grey background shading
  - Tables with styled header row (navy background, white text)
  - Horizontal rules as paragraph borders
  - Page numbers in footer (Page X of Y)
  - Custom colour scheme: navy headings, blue accents

Install: pip install python-docx requests
"""

import re
import os
import tempfile
from typing import List, Dict, Optional

import requests
from core.logger import setup_logger

logger = setup_logger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Exporter
# ─────────────────────────────────────────────────────────────────────────────

class DOCXExporter:
    """
    Converts a markdown research report to a styled .docx file.
    Requires: pip install python-docx requests
    """

    # Brand colours (hex without #)
    TITLE_COLOR   = "1a1a2e"
    HEADING_COLOR = "1e3a5f"
    ACCENT_COLOR  = "2563eb"
    BODY_COLOR    = "1a1a1a"
    REF_COLOR     = "4b5563"
    CODE_BG       = "f3f4f6"
    TABLE_HEADER  = "1e3a5f"

    def export(
        self,
        markdown_report : str,
        output_path     : str,
        title           : str        = "Research Report",
        images          : List[Dict] = None,
    ) -> str:
        """
        Converts markdown to .docx and writes to output_path.
        Returns output_path on success.

        Parameters
        ----------
        markdown_report : Full markdown report string
        output_path     : Where to save the .docx file
        title           : Document title (used in title style)
        images          : Optional list of image dicts for any images
                          not already embedded via markdown image syntax
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "python-docx not installed. Run: pip install python-docx"
            )

        images = images or []
        logger.info(
            f"DOCXExporter | converting | title='{title[:50]}' | "
            f"images={len(images)} | path={output_path}"
        )

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        doc = Document()

        # ── Page setup (A4) ───────────────────────────────────────────────────
        section = doc.sections[0]
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = section.right_margin  = Cm(2.5)
        section.top_margin   = section.bottom_margin = Cm(2.5)

        # ── Apply styles ──────────────────────────────────────────────────────
        self._apply_styles(doc)

        # ── Add page numbers to footer ────────────────────────────────────────
        self._add_page_numbers(doc)

        # ── Parse and render markdown line-by-line ────────────────────────────
        lines = markdown_report.split("\n")
        i     = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # H1 — report title
            if re.match(r"^# (?!#)", line):
                self._add_heading(doc, line[2:].strip(), level=0)

            # H2 — section headings
            elif re.match(r"^## (?!#)", line):
                self._add_heading(doc, line[3:].strip(), level=1)

            # H3 — sub-headings
            elif re.match(r"^### ", line):
                self._add_heading(doc, line[4:].strip(), level=2)

            # Horizontal rule
            elif line.strip() in ("---", "***", "___"):
                self._add_horizontal_rule(doc)

            # Table — collect all consecutive table rows
            elif line.startswith("|") and "|" in line:
                table_lines = []
                while i < len(lines) and lines[i].startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue   # i already advanced

            # Image embed
            elif re.match(r"!\[.*?\]\(.*?\)", line):
                m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
                if m:
                    alt, url = m.group(1), m.group(2)
                    # Check next line for italic caption
                    caption = ""
                    if (
                        i + 1 < len(lines)
                        and lines[i + 1].strip().startswith("*")
                        and lines[i + 1].strip().endswith("*")
                    ):
                        caption = lines[i + 1].strip().strip("*")
                        i += 1
                    self._add_image(doc, url, alt, caption)

            # Bullet list
            elif line.startswith("- ") or line.startswith("* "):
                self._add_list_item(doc, line[2:], numbered=False)

            # Numbered list
            elif re.match(r"^\d+\. ", line):
                text = re.sub(r"^\d+\. ", "", line)
                self._add_list_item(doc, text, numbered=True)

            # Blockquote
            elif line.startswith("> "):
                self._add_blockquote(doc, line[2:])

            # Fenced code block — collect until closing ```
            elif line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block(doc, "\n".join(code_lines))

            # Empty line — skip
            elif line.strip() == "":
                pass

            # Normal paragraph
            else:
                self._add_paragraph(doc, line)

            i += 1

        # ── Save ──────────────────────────────────────────────────────────────
        doc.save(output_path)
        logger.info(f"DOCXExporter | saved | {output_path}")
        return output_path

    # ─────────────────────────────────────────────────────────────────────────
    # Style setup
    # ─────────────────────────────────────────────────────────────────────────

    def _apply_styles(self, doc) -> None:
        """Customises built-in styles to match the report design."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        styles = doc.styles

        # Normal body text
        normal = styles["Normal"]
        normal.font.size                    = Pt(11)
        normal.font.color.rgb               = RGBColor(0x1a, 0x1a, 0x1a)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Heading 1 → ## sections
        h1 = styles["Heading 1"]
        h1.font.size                     = Pt(15)
        h1.font.bold                     = True
        h1.font.color.rgb                = RGBColor(0x1e, 0x3a, 0x5f)
        h1.paragraph_format.space_before = Pt(20)
        h1.paragraph_format.space_after  = Pt(8)

        # Heading 2 → ### sub-sections
        h2 = styles["Heading 2"]
        h2.font.size                     = Pt(12)
        h2.font.bold                     = True
        h2.font.color.rgb                = RGBColor(0x37, 0x41, 0x51)
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after  = Pt(6)

        # Title → # main title
        title_style = styles["Title"]
        title_style.font.size             = Pt(22)
        title_style.font.bold             = True
        title_style.font.color.rgb        = RGBColor(0x1a, 0x1a, 0x2e)
        title_style.paragraph_format.space_after = Pt(14)

    # ─────────────────────────────────────────────────────────────────────────
    # Element renderers
    # ─────────────────────────────────────────────────────────────────────────

    def _add_heading(self, doc, text: str, level: int) -> None:
        """Adds a heading. level 0 = Title, 1 = Heading 1, 2 = Heading 2."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        style_map  = {0: "Title", 1: "Heading 1", 2: "Heading 2"}
        style_name = style_map.get(level, "Heading 2")

        p       = doc.add_paragraph()
        p.style = doc.styles[style_name]
        p.add_run(text)

        # Add blue left-border accent on Heading 1
        if level == 1:
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"),   "single")
            left.set(qn("w:sz"),    "16")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), self.ACCENT_COLOR)
            pBdr.append(left)
            pPr.append(pBdr)

    def _add_paragraph(self, doc, text: str) -> None:
        """Adds a normal paragraph, processing inline markdown."""
        p       = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        self._render_inline(p, text)

    def _render_inline(self, paragraph, text: str) -> None:
        """
        Renders inline markdown into paragraph runs.
        Handles: **bold**, *italic*, `code`, [N] citations (superscript), [text](url).
        """
        from docx.shared import Pt, RGBColor

        pattern = re.compile(
            r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[\d+\]|\[.*?\]\(.*?\))"
        )
        parts = pattern.split(text)

        for part in parts:
            if not part:
                continue

            if part.startswith("**") and part.endswith("**"):
                run      = paragraph.add_run(part[2:-2])
                run.bold = True

            elif part.startswith("*") and part.endswith("*"):
                run        = paragraph.add_run(part[1:-1])
                run.italic = True

            elif part.startswith("`") and part.endswith("`"):
                run                = paragraph.add_run(part[1:-1])
                run.font.name      = "Courier New"
                run.font.size      = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)

            elif re.match(r"^\[(\d+)\]$", part):
                # Citation superscript in blue
                num                  = re.match(r"^\[(\d+)\]$", part).group(1)
                run                  = paragraph.add_run(f"[{num}]")
                run.font.superscript = True
                run.font.size        = Pt(8)
                run.font.color.rgb   = RGBColor(0x25, 0x63, 0xeb)

            elif re.match(r"^\[.*?\]\(.*?\)$", part):
                # Hyperlink text
                m = re.match(r"^\[(.*?)\]\((.*?)\)$", part)
                if m:
                    run                = paragraph.add_run(m.group(1))
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
                    run.underline      = True

            else:
                paragraph.add_run(part)

    def _add_list_item(self, doc, text: str, numbered: bool = False) -> None:
        """Adds a bullet or numbered list item."""
        style = "List Number" if numbered else "List Bullet"
        p     = doc.add_paragraph(style=style)
        self._render_inline(p, text)

    def _add_blockquote(self, doc, text: str) -> None:
        """Adds a blockquote paragraph with left indent and italic text."""
        from docx.shared import Pt, RGBColor, Cm
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_after  = Pt(8)
        run                = p.add_run(text)
        run.italic         = True
        run.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)

    def _add_code_block(self, doc, code: str) -> None:
        """Adds a code block with Courier New and grey background shading."""
        from docx.shared import Pt, RGBColor, Cm
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        p.paragraph_format.space_after  = Pt(8)

        run                = p.add_run(code)
        run.font.name      = "Courier New"
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)

        # Grey background shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  self.CODE_BG)
        pPr.append(shd)

    def _add_table(self, doc, table_lines: List[str]) -> None:
        """Adds a table from markdown table lines."""
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Parse rows — skip separator lines (----)
        rows = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue
            rows.append(cells)

        if not rows:
            return

        max_cols = max(len(r) for r in rows)
        table    = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= max_cols:
                    break
                cell      = table.cell(r_idx, c_idx)
                clean_txt = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                cell.text = clean_txt

                if r_idx == 0:
                    # Style header row
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd   = OxmlElement("w:shd")
                    shd.set(qn("w:val"),   "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"),  self.TABLE_HEADER)
                    tc_pr.append(shd)

        doc.add_paragraph()   # spacing after table

    def _add_image(self, doc, url: str, alt: str, caption: str = "") -> None:
        """
        Downloads image from URL and embeds it centered in the document.
        Falls back to placeholder text on any error.
        """
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            resp = requests.get(url, timeout=10, stream=True)
            resp.raise_for_status()

            content_type = resp.headers.get("Content-Type", "image/jpeg")
            ext = {
                "image/jpeg" : ".jpg",
                "image/png"  : ".png",
                "image/gif"  : ".gif",
                "image/webp" : ".webp",
            }.get(content_type.split(";")[0].strip(), ".jpg")

            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                for chunk in resp.iter_content(8192):
                    tmp.write(chunk)
                tmp_path = tmp.name

            p           = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run         = p.add_run()
            run.add_picture(tmp_path, width=Inches(5.5))

            if caption or alt:
                cap_p           = doc.add_paragraph(caption or alt)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.style     = doc.styles["Normal"]
                for r in cap_p.runs:
                    r.italic         = True
                    r.font.size      = Pt(9.5)
                    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            os.unlink(tmp_path)
            logger.debug(f"DOCXExporter | image embedded | {url[:60]}")

        except Exception as e:
            logger.warning(f"DOCXExporter | image failed: {url[:60]} | {e}")
            p       = doc.add_paragraph(f"[Image: {alt}]")
            p.style = doc.styles["Normal"]

    def _add_horizontal_rule(self, doc) -> None:
        """Adds a thin horizontal border line."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p    = doc.add_paragraph()
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "e5e7eb")
        pBdr.append(bot)
        pPr.append(pBdr)

    # ─────────────────────────────────────────────────────────────────────────
    # Footer — Page X of Y
    # ─────────────────────────────────────────────────────────────────────────

    def _add_page_numbers(self, doc) -> None:
        """Inserts 'Page X of Y' into the document footer."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        section = doc.sections[0]
        footer  = section.footer
        p       = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        def _styled_run(text: str):
            r                = p.add_run(text)
            r.font.size      = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            return r

        def _field_run(field_name: str):
            r     = p.add_run()
            begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.text = field_name
            end   = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
            r._r.append(begin); r._r.append(instr); r._r.append(end)
            r.font.size      = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        _styled_run("Page ")
        _field_run("PAGE")
        _styled_run(" of ")
        _field_run("NUMPAGES")