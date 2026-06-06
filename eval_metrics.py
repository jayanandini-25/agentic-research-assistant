"""
eval_metrics.py  —  Agentic Research Assistant Evaluation
==========================================================
Drop this file into the root of your project (next to main.py).
A .docx evaluation report is auto-generated after EVERY run.

USAGE
-----
  python eval_metrics.py --demo
  python eval_metrics.py --session logs/session_abc123.json
  python eval_metrics.py --query "Diffusion models vs GANs" --url http://localhost:8000

INSTALL DEPENDENCIES (one-time)
---------------------------------
  pip install sentence-transformers requests rich --break-system-packages
  npm install -g docx

AUTO-HOOK INTO YOUR PIPELINE
-----------------------------
  In api/routes.py, the hook is already baked in at the bottom of start_research().
  It calls run_auto_eval(response_dict) automatically.

METRICS OVERVIEW  (6 layers, 24 core metrics)
-----------------------------------------------
  Layer 1  — Planner Agent        : 3 metrics
  Layer 2  — Retrieval            : 3 metrics
  Layer 3  — RAG Pipeline         : 4 metrics
  Layer 4  — Summarization        : 5 metrics
  Layer 5  — Report Quality       : 5 metrics
  Layer 6  — Performance & Cost   : 6 metrics
"""

from __future__ import annotations

import argparse
import json
import math
import os
import re
import statistics
import subprocess
import sys
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ── optional deps ──────────────────────────────────────────────────────────
try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

try:
    from sentence_transformers import SentenceTransformer
    _st_model_cache: Any = None
    HAS_ST = True
except ImportError:
    HAS_ST = False
    _st_model_cache = None

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    from rich.console import Console
    from rich.table import Table
    from rich.panel import Panel
    from rich import box
    HAS_RICH = True
    console = Console()
except ImportError:
    HAS_RICH = False
    console = None


# ═══════════════════════════════════════════════════════════════════════════
# Data structures
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class MetricResult:
    name:   str
    value:  Any
    unit:   str  = ""
    status: str  = "ok"   # ok | warn | fail | skip
    note:   str  = ""


@dataclass
class EvalReport:
    session_id:     str
    query:          str
    timestamp:      str
    planner:        list[MetricResult] = field(default_factory=list)
    retrieval:      list[MetricResult] = field(default_factory=list)
    rag:            list[MetricResult] = field(default_factory=list)
    summarizer:     list[MetricResult] = field(default_factory=list)
    report_quality: list[MetricResult] = field(default_factory=list)
    performance:    list[MetricResult] = field(default_factory=list)

    def all_metrics(self) -> list[MetricResult]:
        return (self.planner + self.retrieval + self.rag +
                self.summarizer + self.report_quality + self.performance)

    def score(self) -> float:
        results = [m for m in self.all_metrics() if m.status != "skip"]
        if not results:
            return 0.0
        weight = {"ok": 1.0, "warn": 0.5, "fail": 0.0}
        return round(100 * sum(weight.get(m.status, 0.5) for m in results) / len(results), 1)


# ═══════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════

def _hi(value: float, warn: float, fail: float) -> str:
    return "ok" if value >= warn else ("warn" if value >= fail else "fail")

def _lo(value: float, warn: float, fail: float) -> str:
    return "ok" if value <= warn else ("warn" if value <= fail else "fail")

def _mean(vals: list[float]) -> float:
    return statistics.mean(vals) if vals else 0.0

def _cosine(a: list[float], b: list[float]) -> float:
    if HAS_NUMPY:
        a_arr, b_arr = np.array(a), np.array(b)
        return float(np.dot(a_arr, b_arr) / (np.linalg.norm(a_arr) * np.linalg.norm(b_arr) + 1e-9))
    dot = sum(x * y for x, y in zip(a, b))
    na  = math.sqrt(sum(x * x for x in a)) + 1e-9
    nb  = math.sqrt(sum(x * x for x in b)) + 1e-9
    return dot / (na * nb)

def _get_st_model():
    global _st_model_cache
    if _st_model_cache is None and HAS_ST:
        _st_model_cache = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    return _st_model_cache

def _skip(name: str, note: str) -> MetricResult:
    return MetricResult(name=name, value=None, status="skip", note=note)

def _is_comparison_query(query: str) -> bool:
    """Detect comparison queries — they naturally have lower query-report alignment."""
    comparison_markers = [" vs ", " versus ", " compared to ", " or ", " and ", " vs. "]
    q_lower = query.lower()
    return any(m in q_lower for m in comparison_markers)


# ═══════════════════════════════════════════════════════════════════════════
# LAYER 1  —  Planner Agent
# ═══════════════════════════════════════════════════════════════════════════

def eval_planner(data: dict) -> list[MetricResult]:
    """
    3 metrics — is the planner generating good, diverse, well-scoped questions?
      1. Sub-question count           — healthy range: 5–12
      2. Research dimension coverage  — % of research dimensions covered
      3. Sub-question diversity       — pairwise cosine similarity (lower = more diverse)
    """
    results   = []
    plog      = data.get("planner_log", {})
    questions = data.get("sub_questions", [])

    # 1. Sub-question count
    n_q = len(questions)
    results.append(MetricResult(
        name="Sub-question count",
        value=n_q,
        unit="questions generated",
        status="ok" if 5 <= n_q <= 12 else "warn",
        note="Too few — research will be shallow" if n_q < 5 else
             "Too many — latency & overlap risk" if n_q > 12 else "",
    ))

    # 2. Dimension coverage
    dims_covered = plog.get("dims_covered", 0)
    total_dims   = plog.get("total_dims", 15)
    coverage     = round(dims_covered / max(total_dims, 1), 3)
    critical_gaps = plog.get("critical_gaps", [])
    results.append(MetricResult(
        name="Dimension coverage",
        value=coverage,
        unit=f"({dims_covered}/{total_dims} dims covered)",
        status=_hi(coverage, warn=0.75, fail=0.50),
        note=("Critical gaps: " + ", ".join(critical_gaps)) if critical_gaps else "All dimensions covered",
    ))

    # 3. Sub-question diversity (semantic similarity)
    if HAS_ST and len(questions) >= 2:
        try:
            model = _get_st_model()
            embs  = model.encode(questions, show_progress_bar=False)
            sims  = [
                _cosine(embs[i].tolist(), embs[j].tolist())
                for i in range(len(embs))
                for j in range(i + 1, len(embs))
            ]
            mean_sim = round(_mean(sims), 3)
            results.append(MetricResult(
                name="Sub-question diversity",
                value=mean_sim,
                unit="mean pairwise cosine (lower = more diverse)",
                status=_lo(mean_sim, warn=0.50, fail=0.65),
                note="High overlap — planner is generating near-duplicate questions" if mean_sim > 0.50 else "Questions are well-diversified",
            ))
        except Exception as exc:
            results.append(_skip("Sub-question diversity", f"Embedding error: {exc}"))
    else:
        results.append(_skip("Sub-question diversity", "pip install sentence-transformers  OR  need ≥2 questions"))

    return results


# ═══════════════════════════════════════════════════════════════════════════
# LAYER 2  —  Retrieval
# ═══════════════════════════════════════════════════════════════════════════

def eval_retrieval(data: dict) -> list[MetricResult]:
    """
    3 metrics — did we get enough diverse evidence?
      1. Total docs retrieved    — volume
      2. Source diversity        — active sources & dominant source share
      3. Source failure rate     — worst-performing source
    """
    results            = []
    rlog               = data.get("retrieval_log", {})
    source_counts      = rlog.get("source_counts", {})
    source_failures    = rlog.get("source_failures", {})
    source_total_calls = rlog.get("source_total_calls", {})
    total_docs         = sum(source_counts.values()) if source_counts else data.get("total_docs", 0)

    # 1. Total docs
    results.append(MetricResult(
        name="Total docs retrieved",
        value=total_docs,
        unit="documents",
        status=_hi(total_docs, warn=40, fail=20),
        note="Sparse evidence base — check API keys / source timeouts" if total_docs < 40 else "",
    ))

    # 2. Source diversity
    if source_counts:
        active  = len([s for s, c in source_counts.items() if c > 0])
        max_src = max(source_counts, key=source_counts.get)
        dom_pct = round(source_counts[max_src] / max(total_docs, 1), 3)
        results.append(MetricResult(
            name="Source diversity",
            value=active,
            unit=f"active sources  (dominant: {max_src} = {round(dom_pct*100)}%)",
            status="ok" if dom_pct < 0.45 and active >= 4 else "warn",
            note=f"{max_src} supplies {round(dom_pct*100)}% of docs — bias risk" if dom_pct >= 0.45 else
                 "Too few active sources — check API keys" if active < 4 else "",
        ))
    else:
        results.append(_skip("Source diversity", "Add source_counts to retrieval_log"))

    # 3. Worst-source failure rate
    if source_failures and source_total_calls:
        rates   = {
            src: round(fails / max(source_total_calls.get(src, fails), 1), 3)
            for src, fails in source_failures.items()
        }
        worst   = max(rates, key=rates.get)
        worst_r = rates[worst]
        results.append(MetricResult(
            name="Worst-source failure rate",
            value=worst_r,
            unit=f"({worst}: {round(worst_r*100)}% of calls failed)",
            status=_lo(worst_r, warn=0.20, fail=0.40),
            note="  |  ".join(f"{s}={round(r*100)}%" for s, r in sorted(rates.items(), key=lambda x: -x[1])),
        ))
    else:
        results.append(_skip("Worst-source failure rate", "Add source_failures + source_total_calls to retrieval_log"))

    return results


# ═══════════════════════════════════════════════════════════════════════════
# LAYER 3  —  RAG Pipeline
# ═══════════════════════════════════════════════════════════════════════════

def eval_rag(data: dict) -> list[MetricResult]:
    """
    4 metrics — is the RAG filter keeping the right chunks?
      1. Chunk retention rate     — fraction of chunks passing threshold (0.30–0.80 is healthy)
      2. Mean chunk similarity    — average relevance of kept chunks (excludes fallback chunks)
      3. Weak-RAG questions       — sub-questions whose chunks score poorly (< 0.45 mean)
      4. Fallback chunk count     — how many chunks used the fallback threshold
    """
    results      = []
    rlog         = data.get("rag_log", {})
    total_before = rlog.get("total_chunks_before", 0)
    total_after  = rlog.get("total_chunks_after",  0)
    per_question = rlog.get("per_question", [])

    # 1. Chunk retention rate
    if total_before > 0:
        retention = round(total_after / total_before, 3)
        results.append(MetricResult(
            name="Chunk retention rate",
            value=retention,
            unit=f"({total_after}/{total_before} chunks kept)",
            status="ok" if 0.30 <= retention <= 0.80 else "warn",
            note="Threshold too strict — discarding good chunks" if retention < 0.30 else
                 "Threshold too loose — noise passing through" if retention > 0.80 else "",
        ))
    else:
        results.append(_skip("Chunk retention rate", "Add total_chunks_before/after to rag_log"))

    # 2. Mean chunk similarity (exclude fallback chunks from mean)
    all_scores      = [s for q in per_question for s in q.get("scores", [])]
    fallback_scores = [s for q in per_question for s in q.get("fallback_scores", [])]
    primary_scores  = [s for s in all_scores if s not in fallback_scores] if fallback_scores else all_scores
    if primary_scores:
        mean_sc   = round(_mean(primary_scores), 3)
        unit_note = f"cosine  (min={round(min(primary_scores),3)}  max={round(max(primary_scores),3)})"
        if fallback_scores:
            unit_note += f"  (excl. {len(fallback_scores)} fallback chunks)"
        results.append(MetricResult(
            name="Mean chunk similarity",
            value=mean_sc,
            unit=unit_note,
            status=_hi(mean_sc, warn=0.50, fail=0.40),
            note="Low mean — retrieval not well-aligned with sub-questions" if mean_sc < 0.50 else "",
        ))
    elif all_scores:
        mean_sc = round(_mean(all_scores), 3)
        results.append(MetricResult(
            name="Mean chunk similarity",
            value=mean_sc,
            unit=f"cosine  (ALL {len(all_scores)} chunks used fallback threshold)",
            status="warn",
            note="All chunks used fallback threshold — retrieval was very weak",
        ))
    else:
        results.append(_skip("Mean chunk similarity", "Add scores lists to rag_log.per_question entries"))

    # 3. Weak-RAG questions
    if per_question:
        weak = [
            q.get("question", "?")[:60]
            for q in per_question
            if q.get("scores") and _mean(q["scores"]) < 0.45
        ]
        results.append(MetricResult(
            name="Weak-RAG questions",
            value=len(weak),
            unit=f"of {len(per_question)} sub-questions have poor chunk alignment",
            status="ok" if not weak else ("warn" if len(weak) <= 2 else "fail"),
            note=("Poorly served: " + " | ".join(weak)) if weak else "All sub-questions have solid chunk alignment",
        ))

    # 4. Fallback chunk count
    total_fallback = sum(q.get("fallback_count", 0) for q in per_question)
    if total_after > 0:
        fb_pct = round(total_fallback / max(total_after, 1) * 100, 1)
        results.append(MetricResult(
            name="Fallback chunk count",
            value=total_fallback,
            unit=f"of {total_after} total  ({fb_pct}% used fallback threshold)",
            status="ok" if fb_pct < 15 else ("warn" if fb_pct < 40 else "fail"),
            note="High fallback rate — retrieval or embedding quality is poor" if fb_pct >= 15 else "",
        ))

    return results


# ═══════════════════════════════════════════════════════════════════════════
# LAYER 4  —  Summarization Agent
# ═══════════════════════════════════════════════════════════════════════════

def eval_summarizer(data: dict) -> list[MetricResult]:
    """
    5 metrics — are summaries deep, grounded, and attributed?
      1. Avg self-reported quality  — agent's own quality score (0–1)
      2. Summary retry rate         — fraction that needed initial retries
      3. Verifier re-retrieval rate — fraction re-retrieved by verifier
      4. Source attribution rate    — fraction citing sources
      5. LLM-as-judge quality       — external LLaMA / GPT scoring (1–5)

    FIX E1: reads s.get("summary") — the key the summarizer agent actually outputs.
             Old code used s.get("answer") which always returned None → judge scored
             empty strings → always got 2.33/5.
    FIX E3: was_re_retrieved is now also checked on "verified_summaries" list if
             "summaries" doesn't carry the flag (some pipeline versions store it there).
    """
    results   = []

    # FIX E3: support both "summaries" and "verified_summaries" keys
    summaries = data.get("summaries", [])
    if not summaries:
        summaries = data.get("verified_summaries", [])
    if not summaries:
        return [_skip("Summarizer", "No summaries in session data")]

    # 1. Self-reported quality
    q_scores = [s["quality"] for s in summaries if "quality" in s]
    if q_scores:
        avg_q = round(_mean(q_scores), 3)
        results.append(MetricResult(
            name="Avg self-reported quality",
            value=avg_q,
            unit=f"over {len(q_scores)} summaries  (0–1)",
            status=_hi(avg_q, warn=0.80, fail=0.60),
            note="Low — cross-check with chunk similarity scores",
        ))

    # 2. Initial retry rate (summarizer's own retries)
    retried    = sum(1 for s in summaries if s.get("attempts", 1) > 1)
    retry_rate = round(retried / len(summaries), 3)
    results.append(MetricResult(
        name="Summary retry rate",
        value=retry_rate,
        unit=f"({retried}/{len(summaries)} summaries needed initial retries)",
        status=_lo(retry_rate, warn=0.10, fail=0.30),
        note="High retry rate — poor chunk quality or overly strict min_quality threshold" if retry_rate > 0.10 else "",
    ))

    # 3. Verifier re-retrieval rate (separate from initial retries)
    # FIX E3: check both direct flag and pipeline-level re_retrieved_questions list
    re_retrieved_questions = set(data.get("re_retrieved_questions", []))
    re_retrieved = sum(
        1 for s in summaries
        if s.get("was_re_retrieved", False)
        or s.get("question", "") in re_retrieved_questions
        or s.get("question", "")[:60] in re_retrieved_questions
    )
    re_retrieval_rate = round(re_retrieved / len(summaries), 3)
    results.append(MetricResult(
        name="Verifier re-retrieval rate",
        value=re_retrieval_rate,
        unit=f"({re_retrieved}/{len(summaries)} summaries re-retrieved by verifier)",
        status=_lo(re_retrieval_rate, warn=0.15, fail=0.35),
        note="High re-retrieval — verifier found weak/contradictory answers frequently" if re_retrieval_rate > 0.15 else "",
    ))

    # 4. Source attribution rate
    sourced   = sum(1 for s in summaries if s.get("sources") and len(s["sources"]) > 0)
    attr_rate = round(sourced / len(summaries), 3)
    results.append(MetricResult(
        name="Source attribution rate",
        value=attr_rate,
        unit=f"({sourced}/{len(summaries)} summaries cite sources)",
        status=_hi(attr_rate, warn=0.90, fail=0.70),
        note="Summaries without sources = untraceable claims, hallucination risk" if attr_rate < 0.90 else "",
    ))

    # 5. LLM-as-judge
    judge_score = _llm_judge_summaries(data, summaries)
    if judge_score is not None and judge_score != "LLM_ERROR":
        results.append(MetricResult(
            name="LLM-as-judge quality",
            value=judge_score,
            unit="/ 5.0  (depth + grounding + completeness)",
            status=_hi(judge_score, warn=3.5, fail=2.5),
            note="",
        ))
    elif judge_score == "LLM_ERROR":
        results.append(MetricResult(
            name="LLM-as-judge quality",
            value=None,
            status="fail",
            note="LLM judge call failed (timeout/bad response) — check API connectivity",
        ))
    else:
        results.append(_skip("LLM-as-judge quality", "Set CEREBRAS_API_KEY or OPENAI_API_KEY env var to enable"))

    return results


# ═══════════════════════════════════════════════════════════════════════════
# LAYER 5  —  Report Quality
# ═══════════════════════════════════════════════════════════════════════════

def eval_report_quality(data: dict) -> list[MetricResult]:
    """
    5 metrics — is the final report comprehensive, grounded, and well-structured?
      1. Report word count            — length proxy for depth
      2. Citation count               — [N] inline markers
      3. Source utilisation rate      — how many retrieved docs actually cited
      4. Query-to-report alignment    — semantic similarity: does report answer the query?
      5. LLM-as-judge report score    — factual depth + consistency + grounding + coverage

    FIX E4: Query-report alignment thresholds relaxed for comparison queries.
             Comparison reports (A vs B) naturally score lower on cosine similarity
             because they discuss both entities, not just the query phrase.
             warn threshold: 0.60 → 0.45 for comparison, 0.55 for regular
             fail threshold: 0.40 → 0.30 for comparison, 0.35 for regular

    FIX E5: Source utilisation warn threshold raised from 0.25 → 0.60.
             If you retrieve 30 sources and cite all 30, utilisation=1.0 is correct.
             Old 0.25 threshold was misleadingly permissive.
    """
    results = []
    report  = data.get("report", "")
    sources = data.get("sources", [])
    query   = data.get("query", "")

    if not report:
        return [_skip("Report quality", "No report string in session data")]

    # 1. Word count
    words = len(report.split())
    results.append(MetricResult(
        name="Report word count",
        value=words,
        unit="words",
        status="ok" if words >= 1000 else ("warn" if words >= 400 else "fail"),
        note="Short report — writer may have been truncated" if words < 1000 else "",
    ))

    # 2. Citation count
    cites = len(re.findall(r"\[\d+\]", report))
    results.append(MetricResult(
        name="Inline citations",
        value=cites,
        unit="[N] markers in report",
        status="ok" if cites >= 10 else ("warn" if cites >= 3 else "fail"),
        note="Few citations — writer not grounding claims in sources" if cites < 10 else "",
    ))

    # 3. Source utilisation rate
    # FIX E5: warn threshold raised from 0.25 → 0.60
    if sources:
        cited_nums  = set(int(m) for m in re.findall(r"\[(\d+)\]", report))
        utilisation = round(len(cited_nums) / max(len(sources), 1), 3)
        results.append(MetricResult(
            name="Source utilisation rate",
            value=utilisation,
            unit=f"({len(cited_nums)}/{len(sources)} retrieved docs cited)",
            status=_hi(utilisation, warn=0.60, fail=0.25),
            note="Most retrieved sources never cited — retrieval may be too broad" if utilisation < 0.60 else "",
        ))
    else:
        results.append(_skip("Source utilisation rate", "Add sources list to session data"))

    # 4. Query-to-report semantic alignment
    # FIX E4: relaxed thresholds for comparison queries
    if HAS_ST and query:
        try:
            model      = _get_st_model()
            q_emb      = model.encode([query], show_progress_bar=False)
            r_emb      = model.encode([" ".join(report.split()[:1000])], show_progress_bar=False)
            alignment  = round(_cosine(q_emb[0].tolist(), r_emb[0].tolist()), 3)
            is_compare = _is_comparison_query(query)
            # Comparison queries discuss multiple entities → lower cosine to short query string
            warn_thr = 0.45 if is_compare else 0.55
            fail_thr = 0.30 if is_compare else 0.35
            compare_note = " (comparison query — threshold relaxed)" if is_compare else ""
            results.append(MetricResult(
                name="Query-report alignment",
                value=alignment,
                unit=f"cosine similarity (query vs report){compare_note}",
                status=_hi(alignment, warn=warn_thr, fail=fail_thr),
                note="Report drifted from query — check writer prompt" if alignment < warn_thr else "Report is well-aligned with the original query",
            ))
        except Exception as exc:
            results.append(_skip("Query-report alignment", f"Embedding error: {exc}"))
    else:
        results.append(_skip("Query-report alignment", "pip install sentence-transformers + need query field"))

    # 5. LLM-as-judge report score
    judge_score = _llm_judge_report(data)
    if judge_score is not None and judge_score != "LLM_ERROR":
        results.append(MetricResult(
            name="LLM-as-judge report score",
            value=judge_score,
            unit="/ 5.0  (factual depth + consistency + grounding + coverage)",
            status=_hi(judge_score, warn=3.5, fail=2.5),
            note="",
        ))
    elif judge_score == "LLM_ERROR":
        results.append(MetricResult(
            name="LLM-as-judge report score",
            value=None,
            status="fail",
            note="LLM judge call failed (timeout/bad response) — check API connectivity",
        ))
    else:
        results.append(_skip("LLM-as-judge report score", "Set CEREBRAS_API_KEY or OPENAI_API_KEY to enable"))

    return results


# ═══════════════════════════════════════════════════════════════════════════
# LAYER 6  —  Performance & Cost
# ═══════════════════════════════════════════════════════════════════════════

def eval_performance(data: dict) -> list[MetricResult]:
    """
    6 metrics — how fast and how expensive?
      1. Total pipeline latency   — end-to-end wall-clock time
      2. Slowest phase            — which phase dominates
      3. Retrieval time share     — retrieval as fraction of total (usually the bottleneck)
      4. Phase breakdown          — all phase times in one row
      5. Total tokens used        — input + output across all LLM calls
      6. Tokens per retrieved doc — efficiency ratio
    """
    results      = []
    timing       = data.get("timing", {})
    token_counts = data.get("token_counts", {})

    # 1. Total latency
    total = timing.get("total", 0)
    if total:
        results.append(MetricResult(
            name="Total pipeline latency",
            value=round(total, 1),
            unit="seconds",
            status="ok" if total < 300 else ("warn" if total < 600 else "fail"),
            note="Target <300s. ArXiv timeouts are the most common bottleneck." if total >= 300 else "",
        ))

    # 2. Slowest phase
    phases      = ["planner", "retrieval", "rag", "summarizer", "verifier", "writer"]
    phase_times = {p: timing[p] for p in phases if timing.get(p)}
    if phase_times and total:
        slowest      = max(phase_times, key=phase_times.get)
        slowest_time = phase_times[slowest]
        slowest_pct  = round(slowest_time / total * 100, 1)
        results.append(MetricResult(
            name="Slowest phase",
            value=f"{slowest}  ({round(slowest_time,1)}s  =  {slowest_pct}% of total)",
            unit="",
            status="warn" if slowest_pct > 75 else "ok",
            note="Reduce ArXiv retries or lower per-source timeout" if slowest == "retrieval" and slowest_pct > 75 else "",
        ))

    # 3. Retrieval time share
    ret_time = timing.get("retrieval", 0)
    if ret_time and total:
        ret_share = round(ret_time / total, 3)
        results.append(MetricResult(
            name="Retrieval time share",
            value=ret_share,
            unit=f"({round(ret_time,1)}s / {round(total,1)}s total)",
            status=_lo(ret_share, warn=0.75, fail=0.88),
            note="Retrieval eating most of runtime — reduce ARXIV_TIMEOUT or max retries" if ret_share > 0.75 else "",
        ))

    # 4. Phase breakdown
    if phase_times:
        breakdown = "  |  ".join(
            f"{p}={round(t,1)}s"
            for p, t in sorted(phase_times.items(), key=lambda x: -x[1])
        )
        results.append(MetricResult(
            name="Phase latency breakdown",
            value=breakdown,
            unit="",
            status="ok",
            note="",
        ))

    # 5. Total tokens
    if token_counts:
        total_in  = sum(v.get("input",  0) for v in token_counts.values())
        total_out = sum(v.get("output", 0) for v in token_counts.values())
        total_tok = total_in + total_out
        results.append(MetricResult(
            name="Total tokens used",
            value=total_tok,
            unit=f"(input={total_in}  output={total_out})",
            status="ok",
            note="",
        ))

        # 6. Tokens per doc
        total_docs = sum(data.get("retrieval_log", {}).get("source_counts", {}).values())
        if total_docs:
            results.append(MetricResult(
                name="Tokens per retrieved doc",
                value=round(total_tok / total_docs, 1),
                unit="tokens/doc",
                status="ok",
                note="",
            ))

    return results


# ═══════════════════════════════════════════════════════════════════════════
# LLM-as-judge helpers
# ═══════════════════════════════════════════════════════════════════════════

def _call_llm(prompt: str, max_tokens: int = 300) -> str | None:
    """Returns None if no API key, 'LLM_ERROR' on call failure."""
    if not HAS_REQUESTS:
        return None
    cerebras_key = os.environ.get("CEREBRAS_API_KEY", "")
    openai_key   = os.environ.get("OPENAI_API_KEY",   "")
    if not cerebras_key and not openai_key:
        return None
    if cerebras_key:
        url     = "https://api.cerebras.ai/v1/chat/completions"
        api_key = cerebras_key
        model   = os.environ.get("CEREBRAS_MODEL", "gpt-oss-120b")
    else:
        base    = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com")
        url     = f"{base}/v1/chat/completions"
        api_key = openai_key
        model   = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model, "max_tokens": max_tokens,
               "messages": [{"role": "user", "content": prompt}]}
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=30)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        return "LLM_ERROR"


def _parse_llm_json(text: str) -> dict | None:
    try:
        return json.loads(re.sub(r"```json|```", "", text).strip())
    except Exception:
        try:
            m = re.search(r"\{.*?\}", text, re.DOTALL)
            return json.loads(m.group()) if m else None
        except Exception:
            return None


def _llm_judge_summaries(data: dict, summaries: list | None = None) -> float | str | None:
    """
    FIX E1 (CRITICAL): reads s.get("summary") not s.get("answer").

    The summarizer agent stores its output under the key "summary".
    The old code used s.get("answer") which returned None for every summary,
    causing the LLM judge to score empty strings → always 2.33/5.

    Also accepts pre-resolved summaries list to avoid double lookup.
    """
    if summaries is None:
        summaries = data.get("summaries", data.get("verified_summaries", []))
    if not summaries:
        return None

    sample      = summaries[:3]
    sample_text = "\n\n".join(
        # FIX E1: use "summary" key, fall back to "answer" for backward compatibility
        f"Q: {s.get('question', '')}\nA: {(s.get('summary') or s.get('answer') or '')[:500]}"
        for s in sample
    )

    # Skip judge call if all sampled summaries are empty (indicates a data key mismatch)
    if not any((s.get("summary") or s.get("answer") or "").strip() for s in sample):
        return _skip("LLM-as-judge quality", "All sampled summaries are empty — check 'summary' key in pipeline output")

    prompt = f"""You are evaluating an AI research assistant's summaries.
Query: {data.get('query', '')}

Sample summaries (first {len(sample)} of {len(summaries)}):
{sample_text}

Rate these summaries 1-5 on each dimension:
- depth: Does the answer go beyond surface-level facts with specific details?
- grounding: Are claims backed by evidence from the sources, not speculation?
- completeness: Does it fully and directly answer the question asked?

Return ONLY valid JSON - no preamble, no backticks:
{{"depth": <1-5>, "grounding": <1-5>, "completeness": <1-5>, "reasoning": "<one sentence>"}}"""

    response = _call_llm(prompt, max_tokens=200)
    if response is None:
        return None
    if response == "LLM_ERROR":
        return "LLM_ERROR"
    parsed = _parse_llm_json(response)
    if not parsed:
        return "LLM_ERROR"
    scores = [parsed.get(k) for k in ("depth", "grounding", "completeness")]
    scores = [s for s in scores if isinstance(s, (int, float)) and 1 <= s <= 5]
    return round(sum(scores) / len(scores), 2) if scores else "LLM_ERROR"


def _llm_judge_report(data: dict) -> float | str | None:
    """Returns float score, 'LLM_ERROR' on call failure, or None if no key."""
    report = data.get("report", "")
    if not report:
        return None
    prompt = f"""You are evaluating a research report generated by an AI system.
Query: {data.get('query', '')}
Sources used: {len(data.get('sources', []))}

Report excerpt (first 2500 chars):
{report[:2500]}

Rate the report 1-5 on each dimension:
- factual_depth: Covers the topic with real depth and specifics, not just surface summaries?
- internal_consistency: No contradictions between sections?
- source_grounding: Claims attributed to sources, not fabricated?
- coverage: All major angles of the query addressed?

Return ONLY valid JSON - no preamble, no backticks:
{{"factual_depth": <1-5>, "internal_consistency": <1-5>, "source_grounding": <1-5>, "coverage": <1-5>, "reasoning": "<one sentence>"}}"""
    response = _call_llm(prompt, max_tokens=250)
    if response is None:
        return None
    if response == "LLM_ERROR":
        return "LLM_ERROR"
    parsed = _parse_llm_json(response)
    if not parsed:
        return "LLM_ERROR"
    keys   = ("factual_depth", "internal_consistency", "source_grounding", "coverage")
    scores = [parsed.get(k) for k in keys]
    scores = [s for s in scores if isinstance(s, (int, float)) and 1 <= s <= 5]
    return round(sum(scores) / len(scores), 2) if scores else "LLM_ERROR"


# ═══════════════════════════════════════════════════════════════════════════
# Runner
# ═══════════════════════════════════════════════════════════════════════════

SECTIONS = [
    ("planner",        "PHASE 2     Planner Agent"),
    ("retrieval",      "PHASE 3/4   Hybrid Retrieval"),
    ("rag",            "PHASE 6     RAG Pipeline"),
    ("summarizer",     "PHASE 8     Summarization Agent"),
    ("report_quality", "PHASE 11    Report Quality"),
    ("performance",    "SYSTEM      Performance & Cost"),
]


class EvalRunner:
    def evaluate(self, session_data: dict) -> EvalReport:
        report = EvalReport(
            session_id=session_data.get("session_id", "unknown"),
            query=session_data.get("query", ""),
            timestamp=time.strftime("%Y-%m-%d %H:%M:%S"),
        )
        print("  [eval] Planner metrics...")
        report.planner        = eval_planner(session_data)
        print("  [eval] Retrieval metrics...")
        report.retrieval      = eval_retrieval(session_data)
        print("  [eval] RAG metrics...")
        report.rag            = eval_rag(session_data)
        print("  [eval] Summarizer metrics...")
        report.summarizer     = eval_summarizer(session_data)
        print("  [eval] Report quality metrics...")
        report.report_quality = eval_report_quality(session_data)
        print("  [eval] Performance metrics...")
        report.performance    = eval_performance(session_data)
        return report

    def evaluate_from_file(self, path: str | Path) -> EvalReport:
        with open(path) as f:
            return self.evaluate(json.load(f))

    def evaluate_from_api(self, query: str, base_url: str = "http://localhost:8000",
                          report_type: str = "research", tone: str = "objective") -> EvalReport:
        if not HAS_REQUESTS:
            raise ImportError("pip install requests")
        payload = {"query": query, "report_type": report_type, "tone": tone}
        print(f"  [eval] POST {base_url}/api/v1/research  (may take several minutes)...")
        resp = requests.post(f"{base_url}/api/v1/research", json=payload, timeout=800)
        resp.raise_for_status()
        return self.evaluate(resp.json())


# ═══════════════════════════════════════════════════════════════════════════
# Console output
# ═══════════════════════════════════════════════════════════════════════════

STATUS_ICON  = {"ok": "OK", "warn": "!!", "fail": "XX", "skip": "--"}
STATUS_COLOR = {"ok": "green", "warn": "yellow", "fail": "red", "skip": "dim"}


def _print_rich(report: EvalReport) -> None:
    score       = report.score()
    score_color = "green" if score >= 75 else ("yellow" if score >= 50 else "red")
    console.print()
    console.print(Panel(
        f"[bold]Session:[/bold]   {report.session_id}\n"
        f"[bold]Query:[/bold]     {report.query}\n"
        f"[bold]Timestamp:[/bold] {report.timestamp}\n"
        f"[bold]Overall score:[/bold]  [{score_color}]{score} / 100[/{score_color}]",
        title="[bold blue]  Agentic Research Assistant — Evaluation Report  [/bold blue]",
        border_style="blue",
    ))
    for attr, title in SECTIONS:
        metrics = getattr(report, attr, [])
        if not metrics:
            continue
        tbl = Table(
            title=title, title_style="bold cyan", title_justify="left",
            box=box.SIMPLE_HEAD, show_header=True, header_style="bold",
            expand=True,
        )
        tbl.add_column("Metric",  style="bold", min_width=32)
        tbl.add_column("Value",   min_width=20)
        tbl.add_column("Status",  min_width=6, justify="center")
        tbl.add_column("Note",    style="dim")
        for m in metrics:
            color   = STATUS_COLOR.get(m.status, "white")
            icon    = STATUS_ICON.get(m.status, "?")
            val_str = str(m.value) if m.value is not None else "—"
            if m.unit:
                val_str += f"  {m.unit}"
            tbl.add_row(m.name, val_str, f"[{color}]{icon}[/{color}]", m.note or "")
        console.print(tbl)
    all_m = report.all_metrics()
    ok   = sum(1 for m in all_m if m.status == "ok")
    warn = sum(1 for m in all_m if m.status == "warn")
    fail = sum(1 for m in all_m if m.status == "fail")
    skip = sum(1 for m in all_m if m.status == "skip")
    console.print(
        f"\n[bold]Summary:[/bold]  "
        f"[green]{ok} ok[/green]  "
        f"[yellow]{warn} warn[/yellow]  "
        f"[red]{fail} fail[/red]  "
        f"[dim]{skip} skipped[/dim]\n"
    )


def _print_plain(report: EvalReport) -> None:
    W   = 76
    SEP = "=" * W
    sep = "-" * W
    print()
    print(SEP)
    print("  Agentic Research Assistant - Evaluation Report")
    print(SEP)
    print(f"  Session   : {report.session_id}")
    print(f"  Query     : {report.query}")
    print(f"  Timestamp : {report.timestamp}")
    print(f"  Score     : {report.score()} / 100")
    print(SEP)
    for attr, title in SECTIONS:
        metrics = getattr(report, attr, [])
        if not metrics:
            continue
        print(f"\n  {title}")
        print(sep)
        for m in metrics:
            icon    = STATUS_ICON.get(m.status, "?")
            val_str = str(m.value) if m.value is not None else "-"
            if m.unit:
                val_str += f"  {m.unit}"
            print(f"  [{icon}]  {m.name:<34}  {val_str}")
            if m.note:
                print(f"           {' '*34}  {m.note}")
    all_m = report.all_metrics()
    ok   = sum(1 for m in all_m if m.status == "ok")
    warn = sum(1 for m in all_m if m.status == "warn")
    fail = sum(1 for m in all_m if m.status == "fail")
    skip = sum(1 for m in all_m if m.status == "skip")
    print()
    print(SEP)
    print(f"  Summary:  {ok} ok   {warn} warn   {fail} fail   {skip} skipped")
    print(SEP)
    print()


def print_report(report: EvalReport) -> None:
    if HAS_RICH:
        _print_rich(report)
    else:
        _print_plain(report)


def save_report_json(report: EvalReport, path: str | None = None) -> str:
    path = path or f"eval_report_{report.session_id[:8]}.json"
    out  = {
        "session_id":    report.session_id,
        "query":         report.query,
        "timestamp":     report.timestamp,
        "overall_score": report.score(),
        "sections":      {
            attr: [
                {"name": m.name, "value": m.value, "unit": m.unit,
                 "status": m.status, "note": m.note}
                for m in getattr(report, attr, [])
            ]
            for attr, _ in SECTIONS
        },
    }
    with open(path, "w") as f:
        json.dump(out, f, indent=2)
    return path


# ═══════════════════════════════════════════════════════════════════════════
# DOCX generation  (pure Python — no Node.js dependency)
# ═══════════════════════════════════════════════════════════════════════════

def _docx_add_page_numbers(doc) -> None:
    """Add 'Page X of Y' to footer using proper OOXML field codes."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        section = doc.sections[0]
        footer  = section.footer
        footer.is_linked_to_previous = False
        p       = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        def _add_field(paragraph, field_code: str):
            run   = paragraph.add_run()
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            run._r.append(begin)

            run2  = paragraph.add_run()
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f" {field_code} "
            run2._r.append(instr)

            run3 = paragraph.add_run()
            end  = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run3._r.append(end)

        r = p.add_run("Page ")
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        _add_field(p, "PAGE")
        r = p.add_run(" of ")
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        _add_field(p, "NUMPAGES")
    except Exception:
        pass


def _docx_score_color(score: float):
    from docx.shared import RGBColor
    if score >= 75: return RGBColor(0x1A, 0x6E, 0x3C)
    if score >= 50: return RGBColor(0x7A, 0x5F, 0x00)
    return RGBColor(0x9B, 0x1C, 0x1C)


_STATUS_RGB = {
    "ok":   (0x1A, 0x6E, 0x3C),
    "warn": (0x7A, 0x5F, 0x00),
    "fail": (0x9B, 0x1C, 0x1C),
    "skip": (0x6B, 0x72, 0x80),
}


def _docx_build_cover(doc, report: "EvalReport") -> None:
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    BLUE_DARK = RGBColor(0x1B, 0x3A, 0x6B)
    BLUE_MED  = RGBColor(0x2E, 0x5F, 0xA3)
    GREY      = RGBColor(0x6B, 0x72, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Agentic Research Assistant")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = BLUE_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Pipeline Evaluation Report")
    r.font.size = Pt(18); r.font.color.rgb = BLUE_MED

    score = report.score()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Overall Health Score")
    r.font.size = Pt(11); r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(f"{score}")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = _docx_score_color(score)
    r = p.add_run(" / 100")
    r.font.size = Pt(18); r.font.color.rgb = GREY

    meta_rows = [
        ("Session ID", report.session_id),
        ("Query",      report.query),
        ("Generated",  report.timestamp),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(meta_rows):
        cell_l = table.cell(i, 0)
        cell_l.text = label
        for r in cell_l.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(10)
        tc_pr = cell_l._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "1B3A6B")
        tc_pr.append(shd)
        for r in cell_l.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell_v = table.cell(i, 1)
        cell_v.text = str(value)
        for r in cell_v.paragraphs[0].runs:
            r.font.size = Pt(10)

    doc.add_paragraph()

    all_m = report.all_metrics()
    counts = {
        "PASSED":  sum(1 for m in all_m if m.status == "ok"),
        "WARN":    sum(1 for m in all_m if m.status == "warn"),
        "FAILED":  sum(1 for m in all_m if m.status == "fail"),
        "SKIPPED": sum(1 for m in all_m if m.status == "skip"),
    }
    count_colors = {
        "PASSED": "E6F4EC", "WARN": "FFFBEB",
        "FAILED": "FEE2E2", "SKIPPED": "F3F4F6",
    }
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Metric Summary")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE_DARK

    ct = doc.add_table(rows=2, cols=4)
    ct.style = "Table Grid"
    for j, (label, count) in enumerate(counts.items()):
        ct.cell(0, j).text = label
        for r in ct.cell(0, j).paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9)
        ct.cell(1, j).text = str(count)
        for r in ct.cell(1, j).paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(18)
        for row_i in (0, 1):
            tc_pr = ct.cell(row_i, j)._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), count_colors[label])
            tc_pr.append(shd)


def _docx_build_details(doc, report: "EvalReport") -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    BLUE_DARK = RGBColor(0x1B, 0x3A, 0x6B)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Detailed Evaluation Results")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE_DARK

    for attr, title in SECTIONS:
        metrics = getattr(report, attr, [])
        if not metrics:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x2E, 0x5F, 0xA3)

        table = doc.add_table(rows=1 + len(metrics), cols=4)
        table.style = "Table Grid"

        headers = ["Metric", "Value", "Status", "Note"]
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            for r in cell.paragraphs[0].runs:
                r.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "1B3A6B")
            tc_pr.append(shd)

        for i, m in enumerate(metrics):
            row_idx = i + 1
            val_str = str(m.value) if m.value is not None else "—"
            if m.unit:
                val_str += f"  {m.unit}"
            icon = STATUS_ICON.get(m.status, "?")
            rgb  = _STATUS_RGB.get(m.status, (0, 0, 0))

            table.cell(row_idx, 0).text = m.name
            for r in table.cell(row_idx, 0).paragraphs[0].runs:
                r.bold = True; r.font.size = Pt(9)

            table.cell(row_idx, 1).text = val_str
            for r in table.cell(row_idx, 1).paragraphs[0].runs:
                r.font.size = Pt(9)

            table.cell(row_idx, 2).text = icon
            for r in table.cell(row_idx, 2).paragraphs[0].runs:
                r.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(*rgb)

            table.cell(row_idx, 3).text = m.note or ""
            for r in table.cell(row_idx, 3).paragraphs[0].runs:
                r.font.size = Pt(8); r.italic = True
                r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            if row_idx % 2 == 0:
                for j in range(4):
                    tc_pr = table.cell(row_idx, j)._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), "F4F6FB")
                    tc_pr.append(shd)

        doc.add_paragraph()


def _docx_build_findings(doc, report: "EvalReport") -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    BLUE_DARK = RGBColor(0x1B, 0x3A, 0x6B)
    all_m = report.all_metrics()
    fails = [m for m in all_m if m.status == "fail"]
    warns = [m for m in all_m if m.status == "warn"]

    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Key Findings & Action Items")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE_DARK

    def _findings_table(items, heading, bg_color):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(10)
        r2 = p2.add_run(heading)
        r2.bold = True; r2.font.size = Pt(11)

        t = doc.add_table(rows=1 + len(items), cols=2)
        t.style = "Table Grid"
        for j, h in enumerate(["Metric", "Action Required"]):
            cell = t.cell(0, j)
            cell.text = h
            for rr in cell.paragraphs[0].runs:
                rr.bold = True; rr.font.size = Pt(9)
                rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "1B3A6B")
            tc_pr.append(shd)
        for i, m in enumerate(items):
            t.cell(i + 1, 0).text = m.name
            for rr in t.cell(i + 1, 0).paragraphs[0].runs:
                rr.bold = True; rr.font.size = Pt(9)
            note = m.note or f"{m.name} scored below threshold (value: {m.value})"
            t.cell(i + 1, 1).text = note
            for rr in t.cell(i + 1, 1).paragraphs[0].runs:
                rr.font.size = Pt(9)
            for j in range(2):
                tc_pr = t.cell(i + 1, j)._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), bg_color)
                tc_pr.append(shd)
        doc.add_paragraph()

    if fails:
        _findings_table(fails, "Critical Issues (must fix)", "FEE2E2")
    if warns:
        _findings_table(warns, "Warnings (should fix)", "FFFBEB")
    if not fails and not warns:
        p3 = doc.add_paragraph()
        r3 = p3.add_run("✓  All metrics passed. No critical issues or warnings detected.")
        r3.bold = True; r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0x1A, 0x6E, 0x3C)


def generate_docx_report(report: EvalReport, docx_path: str | None = None) -> str | None:
    """
    Converts EvalReport → formatted .docx using python-docx (pure Python).
    Returns path to .docx, or None if generation failed.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Cm
    except ImportError:
        print("  [eval] WARNING: python-docx not installed — pip install python-docx")
        return None

    this_dir = Path(__file__).parent
    if docx_path is None:
        reports_dir = this_dir / "eval_reports"
        reports_dir.mkdir(exist_ok=True)
        ts        = time.strftime("%Y%m%d_%H%M%S")
        sid       = report.session_id[:8]
        docx_path = str(reports_dir / f"eval_{sid}_{ts}.docx")

    try:
        doc = DocxDocument()

        section = doc.sections[0]
        section.page_width  = Cm(21.6)
        section.page_height = Cm(27.9)
        section.left_margin = section.right_margin  = Cm(2.0)
        section.top_margin  = section.bottom_margin = Cm(2.0)

        _docx_build_cover(doc, report)
        _docx_build_details(doc, report)
        _docx_build_findings(doc, report)
        _docx_add_page_numbers(doc)

        Path(docx_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(docx_path)
        return docx_path

    except Exception as exc:
        print(f"  [eval] DOCX generation error: {exc}")
        import traceback; traceback.print_exc()
        return None


# ═══════════════════════════════════════════════════════════════════════════
# Public hook — called automatically from api/routes.py
# ═══════════════════════════════════════════════════════════════════════════

def run_auto_eval(session_data: dict, docx_path: str | None = None) -> str | None:
    """
    Drop-in hook for api/routes.py.  Runs all eval layers, prints to console,
    and auto-generates a .docx report to eval_reports/.

    Usage in api/routes.py:
        from eval_metrics import run_auto_eval
        run_auto_eval(response_dict)
    """
    print("\n" + "=" * 60)
    print("  EVAL  Running post-pipeline evaluation...")
    print("=" * 60)
    try:
        runner = EvalRunner()
        report = runner.evaluate(session_data)
        print_report(report)
        saved_json = save_report_json(report)
        print("  [eval] Generating DOCX report...")
        docx_out = generate_docx_report(report, docx_path)
        if docx_out:
            print(f"\n  ✅  DOCX evaluation report  →  {docx_out}\n")
        else:
            print(f"\n  📄  JSON evaluation report  →  {saved_json}\n")
        return docx_out
    except Exception as exc:
        print(f"  [eval] ERROR — evaluation failed: {exc}")
        import traceback; traceback.print_exc()
        return None


def _run_and_report(report: EvalReport, json_path: str | None = None,
                    docx_path: str | None = None) -> None:
    print_report(report)
    saved_json = save_report_json(report, json_path)
    print("  [eval] Generating DOCX report...")
    docx_out = generate_docx_report(report, docx_path)
    if docx_out:
        print(f"\n  ✅  DOCX evaluation report  →  {docx_out}\n")
    else:
        print(f"\n  📄  JSON evaluation report  →  {saved_json}\n")


# ═══════════════════════════════════════════════════════════════════════════
# Demo data (mirrors a real session)
# ═══════════════════════════════════════════════════════════════════════════

def _demo_session() -> dict:
    return {
        "session_id": "88f12bd2-524c-4baf-8307-cc04ba9e2f33",
        "query":      "Diffusion models vs GANs",
        "sub_questions": [
            "How are Diffusion models and GANs related?",
            "How do diffusion models compare to GANs in image quality benchmarks?",
            "Key differences in scalability between diffusion models and GANs?",
            "How does inference speed of GANs compare to diffusion models?",
            "What are the main tradeoffs between Diffusion models and GANs?",
            "When should Diffusion models be preferred over GANs?",
            "Strengths of Diffusion models?",
            "Limitations of Diffusion models?",
            "Strengths of Generative Adversarial Networks?",
            "Real-world use cases for diffusion models and GANs?",
        ],
        "planner_log": {
            "dims_covered":    12,
            "total_dims":      15,
            "iterations_used":  2,
            "max_iterations":   2,
            "critical_gaps":   ["recent_advances"],
        },
        "retrieval_log": {
            "source_counts": {
                "semantic_scholar": 32, "tavily": 25, "nature": 24,
                "arxiv": 10, "hackernews": 4, "wikipedia": 2,
            },
            "source_failures":    {"arxiv": 7, "tavily": 1},
            "source_total_calls": {
                "arxiv": 12, "tavily": 12, "semantic_scholar": 12,
                "nature": 12, "hackernews": 10, "wikipedia": 12,
            },
        },
        "rag_log": {
            "total_chunks_before": 119,
            "total_chunks_after":   61,
            "per_question": [
                {"question": "How are Diffusion models and GANs related?",
                 "scores": [0.691, 0.633, 0.688, 0.61, 0.58, 0.528]},
                {"question": "Image quality benchmarks",
                 "scores": [0.695, 0.677, 0.668, 0.685, 0.645, 0.667]},
                {"question": "Key differences in scalability",
                 "scores": [0.708, 0.706, 0.628, 0.665]},
                {"question": "Inference speed comparison",
                 "scores": [0.592, 0.586, 0.583, 0.461, 0.459, 0.562]},
                {"question": "Key similarities",
                 "scores": [0.58, 0.447, 0.399, 0.431, 0.415, 0.418]},
                {"question": "Main tradeoffs",
                 "scores": [0.798, 0.731, 0.639, 0.612]},
                {"question": "When to prefer which",
                 "scores": [0.719, 0.589, 0.566, 0.553]},
                {"question": "Strengths of Diffusion models",
                 "scores": [0.649, 0.627, 0.599, 0.593, 0.554, 0.459]},
                {"question": "Limitations of Diffusion models",
                 "scores": [0.581, 0.457, 0.386, 0.449]},
                {"question": "Real-world use cases",
                 "scores": [0.593, 0.439, 0.357, 0.373, 0.325]},
            ],
        },
        # FIX E1: demo summaries now use "summary" key (matching real pipeline output)
        "summaries": [
            {
                "question": f"Sub-question {i+1}",
                "summary":  "Diffusion models and GANs are both deep generative models. "
                            "Diffusion models work by reversing a Gaussian noising process, "
                            "achieving FID scores of 2.97 on ImageNet-256 (ADM+guidance). "
                            "GANs use adversarial training between generator and discriminator, "
                            "with StyleGAN-XL achieving 2.30 FID but suffering from mode collapse. "
                            "Diffusion models are 10–100× slower at inference but produce more diverse outputs. " * 3,
                "sources":  [f"https://arxiv.org/abs/2021.{1000+j}" for j in range(3)],
                "quality":  1.0,
                "attempts": 1,
                "was_re_retrieved": False,
            }
            for i in range(10)
        ],
        # FIX E3: also expose re_retrieved_questions at pipeline level
        "re_retrieved_questions": [],
        "report": (
            "# Diffusion models vs GANs\n\n"
            "## Executive Summary\n"
            "Diffusion models have emerged as the dominant approach for high-fidelity image generation, "
            "consistently outperforming GANs on standard benchmarks [1][2][3][4][5]. "
            "While GANs remain faster at inference, diffusion models offer superior training stability "
            "and output diversity [6][7][8].\n\n"
            "## Background\n"
            "The landscape of generative models shifted significantly around 2021 [9][10][11]. "
            "GANs, introduced by Goodfellow et al. in 2014, dominated for nearly a decade [12][13]. "
            "Diffusion models reversed this trend [14][15][16].\n\n"
            "## Overview of Diffusion Models\n"
            "A diffusion model learns to reverse a gradual noising process [17][18][19]. "
            "This probabilistic framework provides stable training and high output diversity [20][21].\n\n"
            "## Overview of GANs\n"
            "GANs train a generator network against a discriminator [22][23][24]. "
            "Single-pass inference makes GANs extremely fast but training is notoriously unstable [25][26].\n\n"
            "## Head-to-Head Comparison\n"
            "On FID scores diffusion models consistently rank better [27][28][29][30]. "
            "On inference latency GANs win by orders of magnitude [31][32].\n\n"
            "## Use Cases\n"
            "Choose diffusion for text-to-image, scientific synthesis, and research [38][39][40]. "
            "Choose GANs for real-time video, style transfer, and edge deployment [41][42][43].\n\n"
            "## Conclusion\n"
            "Diffusion models are the better default for quality-first applications [48][49][50]. "
            "GANs retain their niche where inference speed is the primary constraint [51][52].\n"
        ),
        "sources": [{"id": i, "url": f"https://arxiv.org/abs/202{i%4}.{10000+i}"} for i in range(1, 55)],
        "timing": {
            "planner":    5.2,
            "retrieval":  422.0,
            "rag":        47.0,
            "summarizer": 54.0,
            "verifier":   2.5,
            "writer":     65.0,
            "total":      520.0,
        },
        "token_counts": {
            "planner":    {"input":   800, "output":  350},
            "summarizer": {"input": 14000, "output": 4200},
            "verifier":   {"input":  5000, "output":  800},
            "writer":     {"input": 20000, "output": 7000},
        },
    }


# ═══════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Evaluate pipeline session — auto-generates a .docx report",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python eval_metrics.py --demo
  python eval_metrics.py --session logs/session_abc123.json
  python eval_metrics.py --query "Quantum computing overview" --url http://localhost:8000

DOCX report auto-saved to:  eval_reports/eval_<session>_<timestamp>.docx
        """,
    )
    src = parser.add_mutually_exclusive_group()
    src.add_argument("--demo",    action="store_true", help="Run on built-in demo data")
    src.add_argument("--session", type=str,            help="Path to a session JSON file")
    src.add_argument("--query",   type=str,            help="Live query to fire at running server")
    parser.add_argument("--url",      default="http://localhost:8000", help="Server base URL")
    parser.add_argument("--docx-out", default=None,    help="Custom .docx output path")
    parser.add_argument("--json-out", default=None,    help="Custom JSON output path")
    parser.add_argument("--no-docx",  action="store_true", help="Skip DOCX, console output only")

    args   = parser.parse_args()
    runner = EvalRunner()

    def finish(report: EvalReport) -> None:
        if args.no_docx:
            print_report(report)
            saved = save_report_json(report, args.json_out)
            print(f"\n  📄  JSON report saved → {saved}\n")
        else:
            _run_and_report(report, json_path=args.json_out, docx_path=args.docx_out)

    if args.session:
        print(f"\nEvaluating session: {args.session}")
        finish(runner.evaluate_from_file(args.session))
    elif args.query:
        print(f"\nQuery: '{args.query}'")
        try:
            finish(runner.evaluate_from_api(args.query, base_url=args.url))
        except Exception as exc:
            print(f"Error: {exc}")
            sys.exit(1)
    else:
        print("\nRunning on built-in demo data.")
        print("Usage:  python eval_metrics.py --help\n")
        finish(runner.evaluate(_demo_session()))


if __name__ == "__main__":
    main()